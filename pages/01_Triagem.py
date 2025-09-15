import streamlit as st
import database as db
from datetime import datetime
import pandas as pd
import re
from docx import Document
from io import BytesIO
import locale
import os
import altair as alt # Importa a biblioteca para gráficos

# --- CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(layout="wide", page_title="Triagem e Consulta")

st.title("📋 Triagem e Consulta de Demandas")
st.markdown("Utilize o painel de controlo para visualizar estatísticas e as abas abaixo para as operações.")
st.divider()

# --- NOVA SECÇÃO: PAINEL DE CONTROLE COM GRÁFICOS ---
st.subheader("📊 Painel de Controle")

# Carrega os dados para os gráficos
try:
    df_demandas_chart = db.consultar_demandas()
    df_analises_chart = db.consultar_analises()

    if df_demandas_chart.empty and df_analises_chart.empty:
        st.info("Ainda não há dados suficientes para exibir os gráficos. Comece por registrar demandas ou análises.")
    else:
        col1, col2, col3 = st.columns(3)

        # Gráfico 1: Demandas por Defensor
        with col1:
            if not df_demandas_chart.empty:
                st.markdown("##### Demandas por Defensor(a)")
                chart_defensor = alt.Chart(df_demandas_chart).mark_bar().encode(
                    x=alt.X('defensor', sort='-y', title="Defensor(a)"),
                    y=alt.Y('count()', title="Nº de Demandas"),
                    tooltip=['defensor', 'count()']
                ).interactive()
                st.altair_chart(chart_defensor, use_container_width=True)
            else:
                st.info("Sem dados de demandas para exibir.")
        
        # Gráfico 2: Status das Demandas
        with col2:
            if not df_demandas_chart.empty:
                st.markdown("##### Status das Demandas")
                chart_status = alt.Chart(df_demandas_chart).mark_arc(innerRadius=50).encode(
                    theta=alt.Theta("count()", type="quantitative"),
                    color=alt.Color(field="status", type="nominal", title="Status"),
                    tooltip=['status', 'count()']
                ).properties(
                    width=250,
                    height=250
                )
                st.altair_chart(chart_status, use_container_width=True)
            else:
                st.info("Sem dados de status para exibir.")

        # Gráfico 3: Resultado e Motivo das Análises
        with col3:
            if not df_analises_chart.empty:
                st.markdown("##### Resultado e Motivo das Análises")
                chart_analises = alt.Chart(df_analises_chart).mark_bar().encode(
                    x=alt.X('resultado:N', title="Resultado", axis=alt.Axis(labelAngle=0)),
                    y=alt.Y('count():Q', title="Nº de Análises"),
                    color=alt.Color('motivo:N', title="Motivo"),
                    tooltip=['resultado', 'motivo', 'count()']
                ).interactive()
                st.altair_chart(chart_analises, use_container_width=True)
            else:
                st.info("Sem dados de análises para exibir.")

except Exception as e:
    st.error(f"Ocorreu um erro ao gerar o painel de controlo: {e}")

st.divider()

# --- FUNÇÕES AUXILIARES ---
def formatar_cpf(cpf):
    """Remove caracteres não numéricos do CPF."""
    if cpf:
        return re.sub(r'[^0-9]', '', cpf)
    return ""

def formatar_cpf_para_exibicao(cpf_numerico):
    """Formata um CPF numérico para o formato 000.000.000-00."""
    if not cpf_numerico or len(str(cpf_numerico)) != 11:
        return "" 
    
    cpf_str = str(cpf_numerico)
    return f"{cpf_str[:3]}.{cpf_str[3:6]}.{cpf_str[6:9]}-{cpf_str[9:]}"

def formatar_data_para_extenso(dt):
    """Formata uma data para o formato 'dd de Mês de yyyy' em português."""
    meses = {
        1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
        5: "maio", 6: "junho", 7: "julho", 8: "agosto",
        9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
    }
    return f"{dt.day} de {meses[dt.month]} de {dt.year}"

# --- FUNÇÕES DE LÓGICA DE NEGÓCIO (ANÁLISE DE HIPOSSUFICIÊNCIA) ---
SALARIO_MINIMO = 1518

def avaliar_pf(respostas):
    """Avalia a hipossuficiência de uma Pessoa Física."""
    renda_individual = float(respostas.get("renda_individual", 0))
    renda_familiar = float(respostas.get("renda_familiar", 0))
    investimentos = respostas.get("investimentos", False)
    socio = respostas.get("socio", False)

    if socio:
        capital_social = float(respostas.get("capital_social", 0))
        qtd_socios = int(respostas.get("qtd_socios", 1))
        if qtd_socios > 0:
            valor_por_socio = capital_social / qtd_socios
            if valor_por_socio > renda_individual or valor_por_socio > renda_familiar:
                return "societario"

    if (renda_individual <= 3 * SALARIO_MINIMO or renda_familiar <= 5 * SALARIO_MINIMO) and not investimentos and not socio:
        return True
    return False

def avaliar_pj_lucrativa(respostas):
    """Avalia a hipossuficiência de uma Pessoa Jurídica com fins lucrativos."""
    if respostas.get("socio_recebe_mais", False) or respostas.get("patrimonio_ultrapassa", False):
        return False

    capital_social = float(respostas.get("capital_social", 0))
    qtd_socios = int(respostas.get("qtd_socios", 1))
    renda_referencia = 5 * SALARIO_MINIMO
    
    if qtd_socios > 0 and (capital_social / qtd_socios > renda_referencia):
        return False
    return True

def avaliar_pj_sem_fins(respostas):
    """Avalia a hipossuficiência de uma Pessoa Jurídica sem fins lucrativos."""
    return respostas.get("atua_hipossuficientes", False)

# --- CRIAÇÃO DAS ABAS ---
tab_hipossuficiencia, tab_registrar, tab_consultar, tab_consultar_analises = st.tabs([
    "⚖️ Análise de Hipossuficiência",
    "Registrar Nova Demanda", 
    "Consultar e Editar Demandas",
    "Consultar Análises Salvas"
])

# --- ABA DE ANÁLISE DE HIPOSSUFICIÊNCIA ---
with tab_hipossuficiencia:
    st.subheader("Análise de Hipossuficiência Econômica")
    st.markdown("Ferramenta para avaliar e **salvar** a análise de enquadramento do assistido nos critérios da DPE/BA.")
    
    dados_analise = {}

    with st.container(border=True):
        st.subheader("Informações Básicas")
        tipo_pessoa = st.radio(
            "Tipo de Pessoa", ["Pessoa Física", "Pessoa Jurídica"],
            horizontal=True, key="hipo_tipo_pessoa"
        )
        dados_analise["tipo_pessoa"] = tipo_pessoa
        
        if tipo_pessoa == "Pessoa Física":
            documento = st.text_input("CPF do Assistido", key="hipo_cpf", help="O CPF será salvo sem formatação.")
            dados_analise["documento"] = formatar_cpf(documento)
        else:
            documento = st.text_input("CNPJ da Empresa", key="hipo_cnpj", help="O CNPJ será salvo sem formatação.")
            dados_analise["documento"] = formatar_cpf(documento) # Reutiliza a função de formatação

    with st.container(border=True):
        st.subheader("Vulnerabilidades (Art. 3º, §1º da Resolução)")
        st.caption("A marcação de qualquer vulnerabilidade dispensa a análise dos critérios econômicos.")
        
        vulnerabilidades = [
            "Crianças e adolescentes", "Pessoas idosas", "Pessoas com deficiência",
            "Mulheres vítimas de violência", "Superendividados ou acidentados de consumo",
            "Discriminação (etnia, cor, gênero, etc.)", "Vítimas de tortura, abuso ou violência",
            "Pessoas LGBT+", "Privados de liberdade", "Populações tradicionais (indígenas, quilombolas, etc.)",
            "Em situação de rua / transtornos mentais / catadores", "Risco iminente à vida ou saúde",
            "Vítimas de graves violações de direitos humanos", "Beneficiários de programas sociais",
            "Vítimas de violência institucional", "Outros grupos vulneráveis"
        ]
        
        col1, col2 = st.columns(2)
        vulnerabilidades_selecionadas = []
        for i, item in enumerate(vulnerabilidades):
            if i % 2 == 0:
                if col1.checkbox(item, key=f"hipo_vuln_{i}"): vulnerabilidades_selecionadas.append(item)
            else:
                if col2.checkbox(item, key=f"hipo_vuln_{i}"): vulnerabilidades_selecionadas.append(item)
        dados_analise["vulnerabilidades"] = "; ".join(vulnerabilidades_selecionadas)

    with st.container(border=True):
        st.subheader("Critérios de Hipossuficiência")
        
        analise_habilitada = not bool(vulnerabilidades_selecionadas)
        if not analise_habilitada:
            st.info("Análise de critérios econômicos dispensada devido à seleção de vulnerabilidade.")

        respostas = {}
        if tipo_pessoa == "Pessoa Física":
            respostas["renda_individual"] = st.number_input("Renda líquida individual (R$)", min_value=0.0, step=100.0, disabled=not analise_habilitada, key="hipo_renda_ind")
            respostas["renda_familiar"] = st.number_input("Renda líquida familiar (R$)", min_value=0.0, step=100.0, disabled=not analise_habilitada, key="hipo_renda_fam")
            respostas["investimentos"] = st.checkbox("Possui investimentos ou patrimônio superior a 20 salários mínimos?", disabled=not analise_habilitada, key="hipo_invest")
            respostas["socio"] = st.checkbox("É sócio de alguma empresa ativa?", disabled=not analise_habilitada, key="hipo_socio")
            
            if respostas["socio"]:
                respostas["capital_social"] = st.number_input("Capital Social da empresa (R$)", min_value=0.0, step=100.0, disabled=not analise_habilitada, key="hipo_capital_pf")
                respostas["qtd_socios"] = st.number_input("Quantidade de Sócios", min_value=1, step=1, disabled=not analise_habilitada, key="hipo_qtd_socios_pf")

        else: # Pessoa Jurídica
            natureza_pj = st.radio(
                "Natureza da Pessoa Jurídica", ["Com fins lucrativos", "Sem fins lucrativos"],
                horizontal=True, disabled=not analise_habilitada, key="hipo_natureza_pj"
            )
            respostas["natureza"] = natureza_pj

            if natureza_pj == "Com fins lucrativos":
                respostas["socio_recebe_mais"] = st.checkbox("Algum sócio recebe mais de 5 salários mínimos?", disabled=not analise_habilitada, key="hipo_socio_recebe")
                respostas["patrimonio_ultrapassa"] = st.checkbox("Patrimônio da empresa ultrapassa 60 salários mínimos?", disabled=not analise_habilitada, key="hipo_patrimonio")
                respostas["capital_social"] = st.number_input("Capital Social da empresa (R$)", min_value=0.0, step=100.0, disabled=not analise_habilitada, key="hipo_capital_pj")
                respostas["qtd_socios"] = st.number_input("Quantidade de Sócios", min_value=1, step=1, disabled=not analise_habilitada, key="hipo_qtd_socios_pj")
            else: # Sem fins lucrativos
                respostas["atua_hipossuficientes"] = st.checkbox("Atua na defesa e garantia de direitos de hipossuficientes?", disabled=not analise_habilitada, key="hipo_atua_hipo")
        
        dados_analise["detalhes"] = str(respostas) # Salva um dict de todos os detalhes económicos

    if st.button("Validar e Salvar Análise", type="primary", key="hipo_validar"):
        resultado_final = ""
        motivo = ""
        
        if not dados_analise.get("documento"):
            st.warning("O campo CPF/CNPJ é obrigatório para salvar a análise.")
        else:
            if vulnerabilidades_selecionadas:
                resultado_final = "Aprovado"
                motivo = "Critério de Vulnerabilidade"
                st.success(f"**{resultado_final.upper()} POR {motivo.upper()}.**")
                st.write(f"Vulnerabilidades identificadas: {', '.join(vulnerabilidades_selecionadas)}.")
            else:
                if tipo_pessoa == "Pessoa Física":
                    resultado = avaliar_pf(respostas)
                    if resultado == "societario":
                        resultado_final = "Negado"
                        motivo = "Valor da cota social ultrapassa a renda declarada"
                        st.error(f"**{resultado_final.upper()}.** {motivo}.")
                    elif resultado:
                        resultado_final = "Aprovado"
                        motivo = "Critério Econômico"
                        st.success(f"**{resultado_final.upper()} POR {motivo.upper()}.**")
                    else:
                        resultado_final = "Negado"
                        motivo = "Renda ultrapassa os limites estabelecidos"
                        st.error(f"**{resultado_final.upper()}.** {motivo} e não há critério de vulnerabilidade.")
                else: # Pessoa Jurídica
                    if respostas["natureza"] == "Com fins lucrativos":
                        resultado = avaliar_pj_lucrativa(respostas)
                        if resultado:
                            resultado_final = "Aprovado"
                            motivo = "Critério para PJ com fins lucrativos"
                            st.success(f"**{resultado_final.upper()}.** A empresa se enquadra nos critérios.")
                        else:
                            resultado_final = "Negado"
                            motivo = "Não atende aos critérios para PJ com fins lucrativos"
                            st.error(f"**{resultado_final.upper()}.** {motivo}.")
                    else: # Sem fins lucrativos
                        resultado = avaliar_pj_sem_fins(respostas)
                        if resultado:
                            resultado_final = "Aprovado"
                            motivo = "Atua na defesa de hipossuficientes"
                            st.success(f"**{resultado_final.upper()}.** A entidade se qualifica como defensora de direitos.")
                        else:
                            resultado_final = "Negado"
                            motivo = "Não atua na defesa de hipossuficientes"
                            st.error(f"**{resultado_final.upper()}.** {motivo}.")
            
            if resultado_final and motivo:
                dados_analise['resultado'] = resultado_final
                dados_analise['motivo'] = motivo
                dados_analise['data_analise'] = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
                try:
                    db.adicionar_analise(**dados_analise)
                    st.toast("Análise salva com sucesso no banco de dados!", icon="💾")
                except Exception as e:
                    st.error(f"Ocorreu um erro ao salvar a análise: {e}")
                    st.warning("Verifique se a função `adicionar_analise` existe e está configurada corretamente no seu ficheiro `database.py`.")

# --- ABA DE REGISTRO ---
with tab_registrar:
    LISTA_SERVIDORES = ["THAIS", "RAYSSA", "WELDER"]
    LISTA_DEFENSORES = [
        'Dra. Ana Carolina 1DP', 'Dr. Caio Cesar 2DP', 'Dr. Matheus Rocha 3DP',
        'Dr. Emerson Halsey 4DP', 'Dr. Matheus Bastos 5DP', 'Dra. Janaína Araújo 6DP',
        'Orientação'
    ]
    DEFENSORES_COM_DEMANDAS_RAPIDAS = ['Dra. Ana Carolina 1DP', 'Dra. Janaína Araújo 6DP']
    LISTA_DEMANDAS_RAPIDAS = [
        'Execução de Alimentos', 'Alimentos', 'Divórcio/RDU', 'Inventário',
        'Alvará', 'Curatela', 'Cível geral', 'Prazos geral'
    ]

    if 'pinned_values' not in st.session_state: st.session_state.pinned_values = {"servidor": None}
    if 'pin_status' not in st.session_state: st.session_state.pin_status = {"servidor": False}
    if 'awaiting_confirmation' not in st.session_state: st.session_state.awaiting_confirmation = False
    if 'form_data_to_save' not in st.session_state: st.session_state.form_data_to_save = {}
    if 'clear_form' not in st.session_state: st.session_state.clear_form = False
    if 'num_processos' not in st.session_state: st.session_state.num_processos = 1
    if 'demanda_desc' not in st.session_state: st.session_state.demanda_desc = ""

    if st.session_state.get('clear_form', False):
        if not st.session_state.pin_status['servidor']: st.session_state.pinned_values['servidor'] = None
        st.session_state.nome_assistido = ""
        st.session_state.codigo = ""
        st.session_state.cpf = ""
        st.session_state.demanda_desc = ""
        for i in range(st.session_state.get('num_processos_old', 1)):
            if f"processo_{i}" in st.session_state: st.session_state[f"processo_{i}"] = ""
        st.session_state.num_processos = 1
        st.session_state.clear_form = False

    def update_pin_status(field):
        is_pinned = st.session_state[f'pin_{field}']
        st.session_state.pin_status[field] = is_pinned
        if is_pinned: st.session_state.pinned_values[field] = st.session_state[field]

    def update_pinned_value(field):
        if st.session_state.pin_status.get(field, False):
            st.session_state.pinned_values[field] = st.session_state[field]

    def adicionar_campo_processo():
        st.session_state.num_processos += 1

    def remover_campo_processo(index_para_remover):
        for i in range(index_para_remover, st.session_state.num_processos - 1):
            st.session_state[f"processo_{i}"] = st.session_state[f"processo_{i+1}"]
        del st.session_state[f"processo_{st.session_state.num_processos - 1}"]
        st.session_state.num_processos -= 1
    
    def buscar_nome_por_cpf():
        cpf_input = st.session_state.get('cpf', '')
        cpf_formatado = formatar_cpf(cpf_input)
        if len(cpf_formatado) == 11:
            df = db.consultar_demandas()
            if not df.empty and 'cpf' in df.columns:
                assistido = df[df['cpf'] == cpf_formatado]
                if not assistido.empty:
                    nome_encontrado = assistido['nome_assistido'].iloc[0]
                    st.session_state.nome_assistido = nome_encontrado

    defensor_selecionado = st.selectbox(
        "Selecione o Defensor(a)", options=LISTA_DEFENSORES, index=None,
        placeholder="Escolha um(a) defensor(a) para exibir o formulário", key="registro_defensor"
    )

    if defensor_selecionado:
        if st.session_state.get('awaiting_confirmation', False):
            st.subheader("Confirmar Registo")
            st.warning("Tem a certeza que deseja guardar esta demanda?")
            col1, col2, _ = st.columns([1, 1, 5])
            with col1:
                if st.button("✔️ Sim, guardar"):
                    data = st.session_state.form_data_to_save
                    try:
                        db.adicionar_demanda(**data)
                        st.toast("Demanda registada com sucesso!", icon="✅")
                        st.session_state.clear_form = True
                        st.session_state.awaiting_confirmation = False
                        st.session_state.form_data_to_save = {}
                        st.rerun()
                    except Exception as e: st.error(f"Ocorreu um erro ao salvar a demanda: {e}")
            with col2:
                if st.button("❌ Não, voltar"):
                    st.session_state.awaiting_confirmation = False
                    st.session_state.form_data_to_save = {}
                    st.rerun()
        else:
            st.subheader(f"Atendimento para: {defensor_selecionado}")

            col1, col2 = st.columns([3, 1])
            with col1:
                servidor_val = st.session_state.pinned_values['servidor']
                servidor_index = LISTA_SERVIDORES.index(servidor_val) if servidor_val in LISTA_SERVIDORES else None
                st.selectbox("Servidor", options=LISTA_SERVIDORES, index=servidor_index, placeholder="Selecione o servidor", key="servidor", on_change=update_pinned_value, args=("servidor",))
            with col2:
                st.write("")
                st.checkbox("Fixar", key="pin_servidor", value=st.session_state.pin_status['servidor'], on_change=update_pin_status, args=("servidor",))

            col_nome, col_cpf, col_cod = st.columns([2,1,1])
            with col_nome: st.text_input("Nome do Assistido", placeholder="Nome completo do assistido", key="nome_assistido")
            with col_cpf: st.text_input("CPF do Assistido", placeholder="000.000.000-00", key="cpf", on_change=buscar_nome_por_cpf, help="Digite o CPF e tecle Enter para buscar o nome.")
            with col_cod: st.text_input("Código de Referência", placeholder="Ex: 12345-67", key="codigo")
            
            st.markdown("**Número do Processo(s)**")
            for i in range(st.session_state.num_processos):
                col_input, col_button = st.columns([0.9, 0.1])
                with col_input: st.text_input(f"Processo {i + 1}", key=f"processo_{i}", label_visibility="collapsed")
                with col_button:
                    if i > 0: st.button("❌", key=f"remover_{i}", on_click=remover_campo_processo, args=(i,))
            st.button("➕ Adicionar processo", on_click=adicionar_campo_processo)
            
            st.divider()

            with st.form(f"form_submit", clear_on_submit=True):
                selecao_demanda_list = []
                if defensor_selecionado in DEFENSORES_COM_DEMANDAS_RAPIDAS:
                    selecao_demanda_list = st.multiselect(
                        "Seleção Rápida de Demanda (Opcional)",
                        options=LISTA_DEMANDAS_RAPIDAS,
                        placeholder="Clique para selecionar uma ou mais demandas"
                    )

                demanda = st.text_area("Descrição da Demanda", height=150, placeholder="Descreva a demanda...", key="demanda_desc")
                
                submitted = st.form_submit_button("✔️ Guardar", use_container_width=True, type="primary")

                if submitted:
                    processos = [st.session_state[f"processo_{i}"] for i in range(st.session_state.num_processos) if st.session_state[f"processo_{i}"].strip()]
                    st.session_state.num_processos_old = st.session_state.num_processos

                    if not st.session_state.servidor or not st.session_state.nome_assistido or not st.session_state.codigo or not demanda:
                        st.warning("Por favor, preencha todos os campos obrigatórios (Servidor, Nome, Código e Demanda).")
                    else:
                        st.session_state.form_data_to_save = {
                            "servidor": st.session_state.servidor, "defensor": defensor_selecionado, "nome_assistido": st.session_state.nome_assistido,
                            "cpf": formatar_cpf(st.session_state.cpf), "codigo": st.session_state.codigo, "demanda": demanda, 
                            "selecao_demanda": "; ".join(selecao_demanda_list), "status": 'Pendente', "data": datetime.now().strftime("%d/%m/%Y"), 
                            "horario": datetime.now().strftime("%H:%M:%S"), "numero_processo": ";".join(processos),
                            "documento_gerado": ""
                        }
                        st.session_state.awaiting_confirmation = True
                        st.rerun()

# --- ABA DE CONSULTA DE DEMANDAS---
with tab_consultar:
    st.subheader("🔍 Ferramenta de Busca e Edição de Demandas")

    if 'doc_generator_open_for' not in st.session_state: st.session_state.doc_generator_open_for = None
    if 'confirming_delete' not in st.session_state: st.session_state.confirming_delete = None

    df_original = db.consultar_demandas()

    if st.session_state.doc_generator_open_for is not None:
        id_demanda = st.session_state.doc_generator_open_for
        demanda_selecionada = df_original[df_original['id'] == id_demanda].iloc[0]

        st.subheader(f"📄 Gerar Documento para: {demanda_selecionada['nome_assistido']}")
        st.caption(f"Registo ID: {id_demanda}")

        if st.button("⬅️ Voltar à Consulta"):
            st.session_state.doc_generator_open_for = None
            st.rerun()

        st.markdown("---")
        tipo_documento = st.selectbox(
            "Tipo de documento:",
            options=["Declaração de comparecimento", "Solicitação de Certidão (CRC)", "Declaração de residência", "Carta convite"],
            index=None, placeholder="Selecione uma opção", key=f"doc_type_{id_demanda}"
        )

        if tipo_documento == "Declaração de comparecimento":
            st.markdown("---")
            st.warning("Atenção: O seu ficheiro modelo .docx deve conter as tags <<horadeinicio>> e <<horafim>>.")
            
            defensor_assinatura = st.text_input("Nome do Defensor(a) para Assinatura", value=demanda_selecionada['defensor'], key=f"defensor_assinatura_{id_demanda}")
            
            col_hora1, col_hora2 = st.columns(2)
            with col_hora1: hora_inicio = st.time_input("Hora de início do atendimento", key=f"hinicio_{id_demanda}")
            with col_hora2: hora_fim = st.time_input("Hora de fim do atendimento", key=f"hfim_{id_demanda}")

            if st.button("Gerar e Salvar Documento", key=f"gerar_dec_comp_{id_demanda}", type="primary"):
                if not hora_inicio or not hora_fim or len(defensor_assinatura.strip()) < 3:
                    st.warning("Por favor, preencha todos os campos (horas e nome do defensor com pelo menos 3 caracteres).")
                else:
                    try:
                        caminho_modelo = os.path.join('modelos', 'ADM - DECLARAÇÃO DE COMPARECIMENTO.docx')
                        doc = Document(caminho_modelo)
                        data_extenso = formatar_data_para_extenso(datetime.now())
                        qualificacao_auto = f"CPF/MF: {formatar_cpf_para_exibicao(demanda_selecionada.get('cpf'))}"

                        substituicoes = {
                            "<<NOME PARA ATESTADO DE COMPARECIMENTO>>": demanda_selecionada['nome_assistido'], "<<qualificação>>": qualificacao_auto,
                            "<<dataatendimento>>": data_extenso, "<<COMARCA>>": "Teixeira de Freitas",
                            "<<nomedefensor>>": defensor_assinatura, "<<horadeinicio>>": hora_inicio.strftime("%H:%M"),
                            "<<horafim>>": hora_fim.strftime("%H:%M")
                        }

                        for p in doc.paragraphs:
                            texto_paragrafo = p.text
                            for key, value in substituicoes.items(): texto_paragrafo = texto_paragrafo.replace(key, str(value))
                            if p.text != texto_paragrafo:
                                p.clear()
                                p.add_run(texto_paragrafo)

                        bio = BytesIO()
                        doc.save(bio)
                        
                        st.download_button(
                            label="✔️ Documento Pronto! Clique para Descarregar.", data=bio.getvalue(),
                            file_name=f"Declaracao_Comparecimento_{demanda_selecionada['nome_assistido']}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                        db.atualizar_demanda(id_demanda, {"documento_gerado": tipo_documento})
                        st.toast("Documento gerado e registo atualizado!", icon="📄")

                    except FileNotFoundError: st.error("Erro: O ficheiro modelo 'ADM - DECLARAÇÃO DE COMPARECIMENTO.docx' não foi encontrado.")
                    except Exception as e: st.error(f"Ocorreu um erro ao gerar o documento: {e}")
        
        elif tipo_documento == "Solicitação de Certidão (CRC)":
            st.markdown("---")
            st.info(f"**Assistido(a):** {demanda_selecionada['nome_assistido']}  \n**CPF:** {formatar_cpf_para_exibicao(demanda_selecionada.get('cpf'))}")

            tipo_certidao = st.selectbox("Tipo de Certidão", options=["Nascimento", "Casamento", "Óbito"], key=f"tipo_certidao_{id_demanda}")
            
            dados_certidao = {}
            if tipo_certidao == "Nascimento":
                dados_certidao['crc_nome_registrado'] = st.text_input("Nome Completo do Registrado(a)", value=demanda_selecionada['nome_assistido'], key=f"crc_nome_nasc_{id_demanda}")
                dados_certidao['crc_data_nascimento'] = st.date_input("Data de Nascimento", key=f"crc_data_nasc_{id_demanda}", format="DD/MM/YYYY", value=None)
                dados_certidao['crc_local_nascimento'] = st.text_input("Cidade de Nascimento", key=f"crc_local_nasc_{id_demanda}")
                dados_certidao['crc_nome_pai'] = st.text_input("Nome do Pai", key=f"crc_pai_{id_demanda}")
                dados_certidao['crc_nome_mae'] = st.text_input("Nome da Mãe", key=f"crc_mae_{id_demanda}")

            elif tipo_certidao == "Casamento":
                dados_certidao['crc_nome_registrado'] = st.text_input("Nome do(a) Cônjuge 1", value=demanda_selecionada['nome_assistido'], key=f"crc_nome_cas1_{id_demanda}")
                dados_certidao['crc_nome_conjuge2'] = st.text_input("Nome do(a) Cônjuge 2", key=f"crc_nome_cas2_{id_demanda}")
                dados_certidao['crc_data_casamento'] = st.date_input("Data do Casamento", key=f"crc_data_cas_{id_demanda}", format="DD/MM/YYYY", value=None)
                dados_certidao['crc_local_casamento'] = st.text_input("Cidade do Casamento", key=f"crc_local_cas_{id_demanda}")

            elif tipo_certidao == "Óbito":
                dados_certidao['crc_nome_registrado'] = st.text_input("Nome Completo do(a) Falecido(a)", value=demanda_selecionada['nome_assistido'], key=f"crc_nome_obito_{id_demanda}")
                dados_certidao['crc_data_obito'] = st.date_input("Data do Óbito", key=f"crc_data_obito_{id_demanda}", format="DD/MM/YYYY", value=None)
                dados_certidao['crc_local_obito'] = st.text_input("Cidade do Óbito", key=f"crc_local_obito_{id_demanda}")
                dados_certidao['crc_filiacao_obito'] = st.text_input("Filiação do(a) Falecido(a)", key=f"crc_filiacao_obito_{id_demanda}")

            dados_certidao['crc_cartorio'] = st.text_input("Cartório de Registro (se souber)", key=f"crc_cartorio_{id_demanda}")
            dados_certidao['crc_finalidade'] = st.text_input("Finalidade da Certidão", value="Para fins de prova em processo judicial", key=f"crc_finalidade_{id_demanda}")
            
            if st.button("✔️ Enviar Solicitação para Coordenação", key=f"enviar_crc_{id_demanda}", type="primary", use_container_width=True):
                try:
                    dados_para_salvar = {
                        "documento_gerado": f"Solicitação de {tipo_certidao} Enviada",
                        "crc_tipo_certidao": tipo_certidao, "crc_status": "Pendente"
                    }
                    for key, value in dados_certidao.items():
                        if isinstance(value, datetime): dados_para_salvar[key] = value.strftime('%d/%m/%Y')
                        else: dados_para_salvar[key] = value
                    
                    db.atualizar_demanda(id_demanda, dados_para_salvar)
                    st.toast("Solicitação enviada para a coordenação!", icon="🚀")
                    st.session_state.doc_generator_open_for = None
                    st.rerun()
                except Exception as e: st.error(f"Ocorreu um erro ao enviar a solicitação: {e}")

    else:
        col1, col2, col3 = st.columns(3)
        with col1: filtro_nome = st.text_input("Buscar por Nome do Assistido")
        with col2: filtro_cpf = st.text_input("Buscar por CPF")
        with col3: filtro_defensor = st.selectbox("Filtrar por Defensor", options=["Todos"] + LISTA_DEFENSORES, key="consulta_defensor")

        df_filtrado = df_original.copy()

        if filtro_nome: df_filtrado = df_filtrado[df_filtrado['nome_assistido'].str.contains(filtro_nome, case=False, na=False)]
        if filtro_cpf: df_filtrado = df_filtrado[df_filtrado['cpf'].str.contains(formatar_cpf(filtro_cpf), case=False, na=False)]
        if filtro_defensor != "Todos": df_filtrado = df_filtrado[df_filtrado['defensor'] == filtro_defensor]

        if df_filtrado.empty:
            st.info("Nenhum registo encontrado com os filtros aplicados.")
        else:
            st.info(f"{len(df_filtrado)} registo(s) encontrado(s). Clique num registo para ver os detalhes, editar ou gerar documentos.")
            
            for index, row in df_filtrado.iterrows():
                id_demanda = row['id']
                
                if st.session_state.confirming_delete == id_demanda:
                    st.warning(f"**Tem a certeza que deseja apagar permanentemente o registo de {row['nome_assistido']}?**")
                    col_confirm, col_cancel, _ = st.columns([1,1,5])
                    with col_confirm:
                        if st.button("✔️ Sim, apagar", key=f"confirm_delete_{id_demanda}", type="primary"):
                            try:
                                db.deletar_demanda(id_demanda)
                                st.toast("Registo apagado com sucesso!", icon="🗑️")
                                st.session_state.confirming_delete = None
                                st.rerun()
                            except Exception as e: st.error(f"Erro ao apagar o registo: {e}")
                    with col_cancel:
                        if st.button("❌ Não, cancelar", key=f"cancel_delete_{id_demanda}"):
                            st.session_state.confirming_delete = None
                            st.rerun()

                col_expander, col_delete_btn = st.columns([0.95, 0.05])

                with col_expander:
                    status_color = "green" if row['status'] == 'Pendente' else 'orange'
                    doc_icon = "📄" if row.get('documento_gerado') else ""
                    with st.expander(f"{doc_icon} **{row['nome_assistido']}** |  Data: {row['data']}  |  Status: :{status_color}[{row['status']}]"):
                        
                        with st.form(key=f"form_edit_{id_demanda}"):
                            st.subheader(f"Editando Registo ID: {id_demanda}")
                            dados_editados = {}
                            
                            c1, c2 = st.columns(2)
                            with c1:
                                dados_editados['nome_assistido'] = st.text_input("Nome", value=row['nome_assistido'], key=f"nome_{id_demanda}")
                                dados_editados['cpf'] = st.text_input("CPF", value=row.get('cpf', ''), key=f"cpf_{id_demanda}")
                                dados_editados['servidor'] = st.selectbox("Servidor", options=LISTA_SERVIDORES, index=LISTA_SERVIDORES.index(row['servidor']) if row['servidor'] in LISTA_SERVIDORES else 0, key=f"servidor_{id_demanda}")
                            with c2:
                                dados_editados['codigo'] = st.text_input("Código de Referência", value=row['codigo'], key=f"codigo_{id_demanda}")
                                dados_editados['defensor'] = st.selectbox("Defensor", options=LISTA_DEFENSORES, index=LISTA_DEFENSORES.index(row['defensor']) if row['defensor'] in LISTA_DEFENSORES else 0, key=f"defensor_{id_demanda}")
                                dados_editados['status'] = st.selectbox("Status", options=['Pendente', 'Lido', 'Resolvido', 'Arquivado'], index=['Pendente', 'Lido', 'Resolvido', 'Arquivado'].index(row['status']) if row['status'] in ['Pendente', 'Lido', 'Resolvido', 'Arquivado'] else 0, key=f"status_{id_demanda}")

                            dados_editados['numero_processo'] = st.text_input("Nº Processo", value=row.get('numero_processo', ''), help="Separar múltiplos com (;)", key=f"processo_{id_demanda}")
                            dados_editados['demanda'] = st.text_area("Descrição da Demanda", value=row['demanda'], height=150, key=f"demanda_{id_demanda}")
                            dados_editados['documento_gerado'] = st.text_input("Documento Gerado", value=row.get('documento_gerado', ''), key=f"doc_gerado_{id_demanda}")

                            col_btn1, col_btn2 = st.columns(2)
                            with col_btn1:
                               if st.form_submit_button("💾 Salvar Alterações", type="primary", use_container_width=True):
                                    try:
                                        dados_editados['cpf'] = formatar_cpf(dados_editados['cpf'])
                                        db.atualizar_demanda(id_demanda, dados_editados)
                                        st.toast(f"Registo de {dados_editados['nome_assistido']} atualizado!", icon="🎉")
                                        st.rerun()
                                    except Exception as e: st.error(f"Erro ao atualizar o registo: {e}")
                            with col_btn2:
                                if st.form_submit_button("📄 Gerar/Enviar Solicitação", use_container_width=True):
                                    st.session_state.doc_generator_open_for = id_demanda
                                    st.rerun()
                
                with col_delete_btn:
                    st.write("") 
                    if st.button("❌", key=f"delete_{id_demanda}", help="Apagar este registo"):
                        st.session_state.confirming_delete = id_demanda
                        st.rerun()

# --- ABA DE CONSULTA DE ANÁLISISES ---
with tab_consultar_analises:
    st.subheader("📋 Consulta de Análises de Hipossuficiência Salvas")
    
    try:
        df_analises = db.consultar_analises()
        
        if df_analises.empty:
            st.info("Nenhuma análise de hipossuficiência foi salva ainda.")
        else:
            col1, col2 = st.columns(2)
            with col1:
                filtro_documento = st.text_input("Buscar por CPF/CNPJ", key="filtro_doc_analise")
            with col2:
                filtro_resultado = st.selectbox("Filtrar por Resultado", options=["Todos", "Aprovado", "Negado"], key="filtro_res_analise")

            df_filtrado = df_analises.copy()
            if filtro_documento:
                df_filtrado = df_filtrado[df_filtrado['documento'].str.contains(formatar_cpf(filtro_documento), na=False)]
            if filtro_resultado != "Todos":
                df_filtrado = df_filtrado[df_filtrado['resultado'] == filtro_resultado]

            st.dataframe(df_filtrado, use_container_width=True)

    except Exception as e:
        st.error(f"Ocorreu um erro ao consultar as análises salvas: {e}")
        st.warning("Verifique se a função `consultar_analises` existe e está configurada corretamente no seu ficheiro `database.py`.")

